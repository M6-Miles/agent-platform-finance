#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成三份 Word 文档（企业级重写版 V2）
  1. docs/项目总结报告.docx
  2. docs/项目小白说明书.docx
  3. docs/米志轩-通用Agent平台及证券金融分析应用-工作总结(7.29-8.5).docx
"""
from __future__ import annotations
import io
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
else:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT         = Path(__file__).resolve().parent.parent
DOCS         = ROOT / "docs"
SUMMARY      = DOCS / "项目总结报告.docx"
BEGINNER     = DOCS / "项目小白说明书.docx"
WORK_SUMMARY = DOCS / "米志轩-通用Agent平台及证券金融分析应用-工作总结(7.29-8.5).docx"
FONT         = "微软雅黑"
BODY_SIZE    = 10.5


def _east_asia_font(rPr, name: str) -> None:
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:hint"),    "eastAsia")
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"),    name)
    rFonts.set(qn("w:hAnsi"),    name)
    rPr.insert(0, rFonts)


def set_run_font(run, name: str = FONT, size: float = BODY_SIZE,
                 bold: bool = False, color=None) -> None:
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    _east_asia_font(run._r.get_or_add_rPr(), name)


def set_cell_shading(cell, fill: str = "D9E1F2") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_cell_text(cell, text: str, bold: bool = False,
                  size: float = BODY_SIZE,
                  align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    set_run_font(run, bold=bold, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width   = Cm(21)
    sec.page_height  = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(1.27))
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BODY_SIZE)
    _east_asia_font(style.element.get_or_add_rPr(), FONT)


def add_title(doc: Document, text: str, level: int = 1,
              centered: bool = False) -> None:
    sizes = {0: 20, 1: 16, 2: 13, 3: 11.5}
    sz = sizes.get(level, 12)
    para = doc.add_paragraph()
    if centered:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_font(run, size=sz, bold=True)


def add_section(doc: Document, text: str, level: int = 2) -> None:
    sizes = {2: 12, 3: 11}
    sz = sizes.get(level, 11)
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=sz, bold=True)


def add_body(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    para = doc.add_paragraph(style="List Bullet")
    indent = "    " * level
    run = para.add_run(indent + text)
    set_run_font(run)


def add_numbered(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    set_run_font(run)


def add_kv_table(doc: Document, rows: list[tuple[str, str]],
                 header_fill: str = "2E74B5") -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        row = table.rows[i]
        k_cell = row.cells[0]
        v_cell = row.cells[1]
        set_cell_margins(k_cell)
        set_cell_margins(v_cell)
        if i == 0:
            set_cell_shading(k_cell, header_fill)
            set_cell_shading(v_cell, header_fill)
            set_cell_text(k_cell, k, bold=True, size=10.5)
            set_cell_text(v_cell, v, bold=True, size=10.5)
        else:
            fill = "EBF0FA" if i % 2 == 0 else "FFFFFF"
            set_cell_shading(k_cell, "D9E1F2")
            set_cell_shading(v_cell, fill)
            set_cell_text(k_cell, k, bold=True, size=10)
            set_cell_text(v_cell, v, bold=False, size=10)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Document 1: 项目总结报告
# ─────────────────────────────────────────────────────────────────────────────
def build_project_summary() -> None:
    doc = Document()
    configure_document(doc)

    add_title(doc, "项目总结报告", level=0, centered=True)
    add_title(doc, "构建通用 Agent 平台及证券金融分析应用", level=1, centered=True)
    add_body(doc, "实习单位：北京东方国信科技股份有限公司  |  时间：2026年7月29日—8月5日  |  实习生：米志轩")
    doc.add_paragraph()

    # ── 一、概述 ──────────────────────────────────────────────────────────────
    add_title(doc, "一、项目概述与战略定位", level=1)
    add_body(doc, "本项目基于2026年新兴的Harness Engineering V2.0范式，构建一套「通用Agent平台」，"
             "并在其上搭载「证券金融分析应用」，实现从数据获取、多维分析、综合研判到模拟交易的全链路自动化。")
    add_body(doc, "战略定位：不是简单调用大模型API，而是在系统工程层面解决AI的可靠性、可控性和可审计性"
             "三大核心问题，让Agent在企业级场景中真正可用、可信、可维护。")
    add_kv_table(doc, [
        ("维度", "说明"),
        ("项目类别", "企业级AI工程平台 + 证券金融分析应用"),
        ("技术路线", "Harness Engineering V2.0 / ReAct / DAG多Agent调度"),
        ("数据来源", "AkShare（A股公开行情及财务数据，无版权风险）"),
        ("输出形态", "FastAPI（11条路由）+ Streamlit原型 + SQLite持久化"),
        ("代码质量", "358个测试用例，零失败，零错误，~2分钟完成"),
    ])

    # ── 二、范式背景 ────────────────────────────────────────────────────────
    add_title(doc, "二、核心范式——Harness Engineering V2.0", level=1)
    add_body(doc, "传统AI开发困境：大模型输出不稳定、无工具调用能力、无审计链路。"
             "2026年业界提出「Agent = Model + Harness」工程范式：将大模型与外部系统的"
             "一切交互封装在一个生命周期框架（Harness）里，使Agent行为可预测、可审计、可熔断。")
    add_body(doc, "Harness生命周期三阶段：Pre-flight（输入校验与护栏拦截）→ Loop/Run（ReAct推理循环"
             "或DAG图调度）→ Post-flight（输出过滤与Trace记录）。")
    add_kv_table(doc, [
        ("Harness九大组件", "职责描述"),
        ("Provider", "封装LLM API（DeepSeek/Anthropic），统一调用接口，隐藏密钥"),
        ("Memory", "对话历史管理，支持窗口截断，防止上下文溢出"),
        ("ToolRegistry", "注册可调用工具（get_price / run_backtest等），Schema声明"),
        ("AgentRuntime", "ReAct循环引擎：Reason→Act→Observe，max_steps=4"),
        ("Guardrails", "五大护栏，Pre-flight拦截 + Post-flight二次过滤"),
        ("CircuitBreaker", "熔断器：max_failures=3，cooldown_s=300，自动重置"),
        ("LangGraph", "主编排引擎：StateGraph + MemorySaver/SqliteSaver，interrupt/resume 支持人工审批"),
        ("Observability", "全链路可观测：P50/P95延迟、成功率、护栏违规次数、per-agent指标"),
        ("TraceRecord", "每次run()的完整追踪，记录retries/违规/耗时，可导出"),
    ])

    # ── 三、五大Guardrail护栏体系 ─────────────────────────────────────────────
    add_title(doc, "三、五大 Guardrail 护栏体系", level=1)
    add_body(doc, "护栏是Harness的安全闸门，每次Agent运行前后均执行。幻觉阻断率100%，误报率0%（实验验证）。")
    add_kv_table(doc, [
        ("护栏名称", "功能描述"),
        ("JSONSchemaValidator",    "校验LLM输出是否符合预定JSON Schema，不符则拒绝"),
        ("SourceAttributionFilter","要求输出包含source字段和免责声明，防止无来源结论"),
        ("RateLimiter",            "限流控制，防止API滥用和超额计费"),
        ("KeywordBlocker",         "拦截含风险关键词（如绝对涨跌承诺）的输出"),
        ("CrossValidator",         "交叉校验多路Agent结论，检测异常分歧"),
    ])
    add_body(doc, "实验结果：对5条预设幻觉样本（无来源/承诺绝对收益/结构不合规）全部拦截（100%），"
             "对3条正常样本零误报（0%），详见 src/agent_platform/core/harness_experiment.py。")

    # ── 四、专业Agent体系 ──────────────────────────────────────────────────
    add_title(doc, "四、六大专业 Agent + 风控体系", level=1)
    add_body(doc, "金融分析任务被分解为六个正交维度，每个Agent专注一个维度，最终由SynthesisAgent汇总研判。")
    add_kv_table(doc, [
        ("Agent名称", "分析维度 / 评分范围"),
        ("TechnicalAgent",     "技术面：MA/EMA/MACD/RSI/BOLL/KDJ指标（±65分）"),
        ("FundamentalAgent",   "基本面：市盈率/市净率/ROE/营收增速等财务指标（±15分）"),
        ("IndustryAgent",      "行业面：行业景气度、政策利好/利空（±10分）"),
        ("MarketRegimeAgent",  "市场风格：牛市/熊市/震荡判断（±15分）"),
        ("SynthesisAgent",     "综合研判：汇总四路评分，输出buy/sell/hold信号"),
        ("TraderAgent",        "交易决策：解析信号，限制单笔仓位<=10%"),
        ("RiskManager",        "风险管理：仓位>10%强制触发HumanApprovalRequired异常"),
    ])
    add_body(doc, "SynthesisAgent评分逻辑：总分range=-90到+90，"
             "confidence = (score + 90) / 180，"
             "score>0=bull/buy，score<0=bear/sell，score~0=hold。")

    # ── 五、LangGraph 主编排引擎 ───────────────────────────────────────────
    add_title(doc, "五、LangGraph 主编排引擎", level=1)
    add_body(doc, "主工作流已迁移至 LangGraph StateGraph。"
             "LangGraph StateGraph 原生支持 interrupt/resume 人工审批、SQLite checkpoint 持久化、"
             "并行扇出（START → 四路并行 → synthesis → trader → risk → harness）。")
    add_kv_table(doc, [
        ("机制", "实现方式"),
        ("并行扇出",   "START 同时触发 technical/fundamental/industry/market_regime 四个节点"),
        ("条件路由",   "synthesis → 低置信度→END / 可交易→trader；trader → HAR→human_approval / OK→risk_manager"),
        ("人工审批",   "interrupt() 暂停图；前端 POST /research/{thread_id}/resume 传入 approve/reject 恢复"),
        ("Checkpoint", "SqliteSaver 持久化（*_lg_checkpoints.db）；MemorySaver 作为开发/测试降级"),
        ("节点计时",   "每个节点通过 trace_entries 记录真实耗时，ObservabilityPanel 直接读取"),
        ("异常策略",   "TypeError/AttributeError/NameError → reraise；其余 Exception → 写入 errors 字段"),
    ])

    # ── 六、回测引擎与数据层 ──────────────────────────────────────────────
    add_title(doc, "六、回测引擎与数据层", level=1)
    add_body(doc, "回测引擎模拟真实交易环境，对历史信号进行绩效评估，是策略可信度的核心验证工具。")
    add_kv_table(doc, [
        ("模块", "说明"),
        ("数据来源",   "AkShare获取A股日线行情（OHLCV），无需付费API"),
        ("技术指标",   "MA/EMA/MACD/RSI/BOLL/KDJ/ATR/CCI，共10类指标函数"),
        ("交易规则",   "出现buy信号则下一根K线开盘买入，出现sell信号则开盘卖出"),
        ("滑点成本",   "0.1%（双边，模拟真实市场冲击）"),
        ("手续费",     "0.03%（双边，符合A股市场实际）"),
        ("绩效指标",   "Sharpe Ratio / Max Drawdown / 年化收益率 / 胜率"),
        ("E-01结果",   "真实股票Sharpe=-0.440（未达标），合成数据Sharpe=+0.33（接近标准）"),
    ])
    add_body(doc, "E-01未达标原因诚实说明：A股历史数据中，纯技术指标信号在随机走势下"
             "难以稳定获得Sharpe>0.5。这是金融领域的客观挑战，"
             "不通过选择有利时段或注入正漂移来人为美化结果（遵循SPEC.md §3.1）。")

    # ── 七、API接口与UI ─────────────────────────────────────────────────
    add_title(doc, "七、API 接口层与用户界面", level=1)
    add_body(doc, "FastAPI提供11条REST路由，Streamlit提供可视化原型，SQLite负责数据持久化。")
    add_kv_table(doc, [
        ("接口路由", "功能"),
        ("POST /auth/register",       "用户注册，密码bcrypt哈希存储"),
        ("POST /auth/login",          "登录，返回JWT Token"),
        ("GET  /stocks/{code}/price", "获取实时/历史行情"),
        ("POST /analysis/run",        "触发多Agent分析流水线"),
        ("GET  /analysis/{id}",       "查询分析任务状态"),
        ("POST /backtest/run",        "启动回测，支持自定义时间段"),
        ("GET  /backtest/{id}",       "查询回测结果（含Sharpe/MaxDD）"),
        ("GET  /agents/status",       "查询所有Agent运行状态"),
        ("GET  /observability/summary","获取全局可观测指标"),
        ("POST /trade/order",         "提交模拟订单（仓位>10%抛HumanApprovalRequired）"),
        ("GET  /health",              "健康检查端点"),
    ])

    # ── 八、工程质量 ────────────────────────────────────────────────────
    add_title(doc, "八、工程质量与安全规范", level=1)
    add_kv_table(doc, [
        ("质量维度", "具体措施"),
        ("测试覆盖", "358个测试用例全部通过，运行时间~2分钟，零失败零跳过"),
        ("密钥管理", "API Key仅存.env，从不硬编码，.gitignore已排除"),
        ("金融合规", "所有分析输出含source字段和免责声明（仅供研究参考，不构成投资建议）"),
        ("风控硬限", "单笔仓位>10%强制抛HumanApprovalRequired异常，无法绕过"),
        ("禁止真实交易", "MockBroker仅本地模拟，禁止引入requests/httpx连接真实券商"),
        ("幻觉防护", "5大Guardrail实验：幻觉阻断率100%，正常样本误报率0%"),
        ("代码组织", "src/agent_platform/分层架构：core/finance/api/ui各司其职"),
    ])

    # ── 九、验收标准 ────────────────────────────────────────────────────
    add_title(doc, "九、验收标准执行情况（SPEC.md）", level=1)
    add_body(doc, "项目共19条验收标准，关键指标执行情况如下：")
    add_kv_table(doc, [
        ("标准", "状态 / 说明"),
        ("358测试全通过",             "通过 — exit=0，0 failed，0 error"),
        ("Guardrail幻觉阻断率100%",   "通过 — harness_experiment.py验证"),
        ("CircuitBreaker熔断",        "通过 — max_failures=3, cooldown=300s"),
        ("E-01: Sharpe>0.5",          "未达标 — 真实股票-0.440，合成+0.33"),
        ("API 11条路由可用",          "通过 — FastAPI OpenAPI文档可验证"),
        ("仓位>10%触发人工审批",      "通过 — HumanApprovalRequired已实现"),
        ("免责声明强制输出",          "通过 — SourceAttributionFilter保障"),
        ("MockBroker本地隔离",        "通过 — 无外部网络连接"),
    ])

    # ── 十、后续展望 ────────────────────────────────────────────────────
    add_title(doc, "十、后续展望与改进方向", level=1)
    add_bullet(doc, "Sharpe优化：引入多因子策略（动量+反转+基本面联合信号），有望改善E-01")
    add_bullet(doc, "实时行情：接入Level-2行情，提升日内交易信号频率与质量")
    add_bullet(doc, "Agent扩展：新增SentimentAgent（舆情分析）和MacroAgent（宏观经济）")
    add_bullet(doc, "分布式部署：将 LangGraph 执行器接入 Celery/Ray，支持水平扩展")
    add_bullet(doc, "模型微调：基于历史信号标注数据，对LLM进行LoRA微调，提升分析精度")

    doc.save(SUMMARY)
    print(f"[OK] 已生成：{SUMMARY}")


# ─────────────────────────────────────────────────────────────────────────────
# Document 2: 项目小白说明书
# ─────────────────────────────────────────────────────────────────────────────
def build_beginner_guide() -> None:
    doc = Document()
    configure_document(doc)

    add_title(doc, "项目小白说明书", level=0, centered=True)
    add_title(doc, "通用 Agent 平台及证券金融分析应用", level=1, centered=True)
    add_body(doc, "这是一份「吃透代码」的详细说明书。看完后，你将理解这个项目是如何一步步构建的，"
             "每一步解决什么问题，涉及哪些技术概念。适合向导师汇报或深入学习。")
    doc.add_paragraph()

    # ── 一、项目是什么 ──────────────────────────────────────────────────
    add_title(doc, "一、这个项目到底是什么？", level=1)
    add_body(doc, "这是一个「让AI可靠分析股票」的完整工程系统。")
    add_body(doc, "核心问题：大模型（如ChatGPT/DeepSeek）很聪明，但它会「瞎说」——没有数据依据、"
             "承诺绝对收益、输出格式乱七八糟。直接用它做金融分析，后果很严重。")
    add_body(doc, "本项目的解决方案：搭建一个「Harness」框架（线束，像汽车引擎的线束一样），"
             "把大模型包裹起来，让它的输入输出都经过严格检查、不能乱调用工具、不能说瞎话。"
             "在这个可靠的框架上，再搭建6个专业Agent分工分析股票，最终汇总成买卖信号。")
    add_body(doc, "技术路线：Harness Engineering V2.0（2026年新范式）+ ReAct推理循环 + "
             "DAG多Agent并行调度 + 回测验证。")

    # ── 二、为什么这么做 ────────────────────────────────────────────────
    add_title(doc, "二、为什么要这么做？（背景与需求）", level=1)
    add_body(doc, "传统AI开发的三大痛点：")
    add_numbered(doc, "不可靠：大模型输出随机性强，同样问题每次回答可能不一样，还会编造数据")
    add_numbered(doc, "不可控：没有工具调用能力（不能查数据库、不能联网），只能靠训练时的知识瞎猜")
    add_numbered(doc, "不可审计：出了问题不知道哪里错了，没有日志、没有trace、无法追溯")
    add_body(doc, "2026年业界提出新范式：Agent = Model + Harness。"
             "把大模型当成「CPU」，Harness当成「主板+总线」，两者配合才能成为可用的系统。"
             "Harness负责：输入检查、工具注册、循环推理、输出过滤、熔断保护、全程追踪。")

    # ── 三、第一步：搭骨架（Harness九大组件）────────────────────────────
    add_title(doc, "三、第一步：搭 Harness 骨架（解决「让 Agent 可靠运行」的需求）", level=1)
    add_body(doc, "这一步的目标：搭建一个空的Harness框架，还没有具体的金融逻辑，"
             "但已经具备完整的生命周期管理能力。")
    add_body(doc, "Harness包含九大组件，每个组件解决一个具体问题：")
    add_kv_table(doc, [
        ("组件名", "解决什么问题"),
        ("Provider",       "问题：不同大模型API格式不统一。解决：封装统一接口，代码只调用Provider，不直接碰LLM API"),
        ("Memory",         "问题：对话历史太长会超限。解决：管理历史记录窗口，自动截断"),
        ("ToolRegistry",   "问题：Agent需要调用外部工具（查数据库、计算指标），但不知道有哪些工具。解决：注册工具清单+Schema"),
        ("AgentRuntime",   "问题：Agent需要「思考→行动→观察」循环迭代。解决：实现ReAct循环引擎，最多4步"),
        ("Guardrails",     "问题：Agent会瞎说话。解决：5个护栏检查输入输出，不合规就拦截"),
        ("CircuitBreaker", "问题：某个Agent反复失败会拖垮整个系统。解决：3次失败后熔断，冷却5分钟"),
        ("LangGraph",      "问题：多个Agent需要并行执行、按依赖顺序调度、支持人工中断恢复。解决：StateGraph并行扇出+interrupt()/Command(resume=...)"),
        ("Observability",  "问题：不知道系统运行情况。解决：记录每次调用的延迟/成功率/违规次数"),
        ("TraceRecord",    "问题：出问题后无法复现。解决：每次run()都记录完整trace，可导出分析"),
    ])
    add_body(doc, "代码位置：src/agent_platform/core/harness.py （AgentHarness类，约180行）")
    add_body(doc, "核心流程：run() → Pre-flight(check_input) → _run_agent() → Post-flight(validate_output) → record_trace()")

    # ── 四、第二步：Loop引擎（ReAct）───────────────────────────────────
    add_title(doc, "四、第二步：实现 Loop 引擎（解决「Agent如何自主推理」的需求）", level=1)
    add_body(doc, "这一步的目标：让单个Agent能够自主地「思考→选择工具→观察结果→再思考」，"
             "形成完整的推理循环，不需要人工干预每一步。")
    add_body(doc, "实现方式：ReAct（Reasoning + Acting）循环，这是2023年Google提出的Agent推理范式：")
    add_numbered(doc, "Reason（思考）：把当前问题和历史对话发给LLM，让它思考下一步该做什么")
    add_numbered(doc, "Act（行动）：LLM输出工具调用指令（如「调用get_price查茅台股价」）")
    add_numbered(doc, "Observe（观察）：执行工具，把结果返回给LLM作为新的上下文")
    add_numbered(doc, "重复以上1-3步，最多4次（max_steps=4），然后输出最终答案")
    add_body(doc, "代码位置：src/agent_platform/core/agent_runtime.py（AgentRuntime类）")
    add_body(doc, "关键设计：max_steps=4防止无限循环；工具调用结果作为Observe写回history；"
             "最终返回AgentRunResult(answer, steps, provider)三元组。")

    # ── 五、第三步：五大Guardrail────────────────────────────────────────
    add_title(doc, "五、第三步：五大 Guardrail 护栏（解决「防止AI瞎说话」的需求）", level=1)
    add_body(doc, "这一步的目标：给Harness加上安全护栏，确保Agent的输入输出都符合规范，"
             "特别是在金融场景中绝对不能出现「幻觉」（编造数据）或「不当承诺」（保证收益）。")
    add_body(doc, "五大护栏，每个都有check_input()和validate_output()两个方法：")
    add_kv_table(doc, [
        ("护栏名", "拦截什么 / 怎么检查"),
        ("JSONSchemaValidator",     "LLM输出的JSON格式不对？→ 对照预定义Schema校验，不符就拒绝"),
        ("SourceAttributionFilter", "输出没有数据来源？→ 检查source字段和免责声明是否存在"),
        ("RateLimiter",             "调用太频繁？→ 统计时间窗口内的调用次数，超限返回429"),
        ("KeywordBlocker",          "包含「保证收益」「一定涨」等风险词？→ 关键词黑名单匹配，命中即拦"),
        ("CrossValidator",          "多个Agent结论严重矛盾？→ 计算置信度分歧，超阈值标记为异常"),
    ])
    add_body(doc, "验证实验：5条预设幻觉样本（无来源/承诺绝对收益/格式不合规）全部被拦截（100%），"
             "3条正常样本零误报（0%）。代码：src/agent_platform/core/harness_experiment.py")

    # ── 六、第四步：LangGraph 主编排引擎 ──────────────────────────────────
    add_title(doc, "六、第四步：LangGraph 主编排引擎（解决「多 Agent 协作」的需求）", level=1)
    add_body(doc, "这一步的目标：当有多个Agent需要协作时（如：先查行情，再分析技术面和基本面，"
             "最后综合研判），需要一个调度引擎来管理它们的执行顺序和并行关系。")
    add_body(doc, "主工作流使用 LangGraph StateGraph，原生支持并行扇出、条件路由、"
             "interrupt/resume 人工审批和 SQLite checkpoint。")
    add_body(doc, "LangGraph 工作流运行流程：")
    add_numbered(doc, "建图：StateGraph 注册节点和条件边，START 并行触发四路分析节点")
    add_numbered(doc, "并行执行：technical/fundamental/industry/market_regime 同时执行，汇合到 synthesis_agent")
    add_numbered(doc, "条件路由：置信度 ≤30% 直接 END（no_trade）；>30% 进入 trader_agent")
    add_numbered(doc, "人工审批：仓位超阈值 → interrupt() 暂停图 → 前端 approve/reject → Command(resume=...) 恢复")
    add_numbered(doc, "Checkpoint：SqliteSaver 持久化到 *_lg_checkpoints.db，支持跨请求 resume")
    add_body(doc, "代码位置：src/agent_platform/finance/securities_graph.py（build_securities_graph / run_securities_analysis）")

    # ── 七、第五步：数据层和技术指标────────────────────────────────────
    add_title(doc, "七、第五步：数据层和技术指标（解决「Agent需要真实数据」的需求）", level=1)
    add_body(doc, "这一步的目标：Agent不能只会「瞎猜」，必须基于真实的股票行情数据和技术指标做分析。")
    add_body(doc, "数据来源：AkShare——一个开源的A股数据接口库，提供日线行情、财务报表、行业分类等数据，无需付费。")
    add_body(doc, "技术指标库：src/agent_platform/finance/indicators.py，实现了10类常用指标：")
    add_bullet(doc, "趋势类：MA（移动平均）、EMA（指数移动平均）", 1)
    add_bullet(doc, "动量类：MACD（异同移动平均）、RSI（相对强弱指数）", 1)
    add_bullet(doc, "波动类：BOLL（布林带）、ATR（真实波幅）", 1)
    add_bullet(doc, "超买超卖类：KDJ（随机指标）、CCI（商品通道指数）", 1)
    add_bullet(doc, "风险类：annualized_volatility（年化波动率）、max_drawdown（最大回撤）", 1)
    add_body(doc, "数据流：AkShare获取原始OHLCV → indicators.py计算指标 → 以DataFrame形式提供给Agent。")

    # ── 八、第六步：六大专业Agent──────────────────────────────────────
    add_title(doc, "八、第六步：六大专业 Agent（解决「分工分析」的需求）", level=1)
    add_body(doc, "这一步的目标：把复杂的股票分析任务分解为多个正交维度，每个Agent专注一个维度，"
             "避免单个Agent负担过重、分析混乱。")
    add_kv_table(doc, [
        ("Agent", "负责什么 / 输出格式"),
        ("TechnicalAgent",     "技术面分析：看K线、指标（MA/MACD/RSI等），判断超买超卖、趋势强弱。评分±65"),
        ("FundamentalAgent",   "基本面分析：看市盈率、ROE、营收增速，判断公司质地。评分±15"),
        ("IndustryAgent",      "行业面分析：看行业景气度、政策利好利空。评分±10"),
        ("MarketRegimeAgent",  "市场风格分析：判断当前是牛市/熊市/震荡市，决定策略激进度。评分±15"),
        ("SynthesisAgent",     "综合研判：汇总上述四路评分（总分范围-90到+90），输出buy/sell/hold信号"),
        ("TraderAgent",        "交易决策：解析信号，生成订单，限制单笔仓位<=10%"),
    ])
    add_body(doc, "代码位置：src/agent_platform/finance/ 下各Agent模块。")
    add_body(doc, "SynthesisAgent评分逻辑：总分score=tech+fund+industry+regime，"
             "confidence=(score+90)/180，score>0→bull/buy，score<0→bear/sell，score≈0→hold。")

    # ── 九、第七步：综合研判与风控─────────────────────────────────────
    add_title(doc, "九、第七步：综合研判与风控（解决「汇总信号、控制风险」的需求）", level=1)
    add_body(doc, "这一步的目标：把六路Agent的分析结果汇总成一个明确的买卖信号，"
             "并加上硬性风控规则，防止单笔交易仓位过大。")
    add_body(doc, "SynthesisAgent工作流程：")
    add_numbered(doc, "收集四路评分：technical_score、fundamental_score、industry_score、regime_score")
    add_numbered(doc, "计算总分：total_score = sum(四路评分)，范围-90到+90")
    add_numbered(doc, "生成信号：score>10→buy，score<-10→sell，-10<=score<=10→hold")
    add_numbered(doc, "置信度：confidence = (score + 90) / 180，取值0.0到1.0")
    add_numbered(doc, "输出结构：SynthesisResult(signal, confidence, bull_arguments, bear_arguments)")
    add_body(doc, "风控机制（RiskManager + TraderAgent）：")
    add_bullet(doc, "单笔仓位硬限：position > 10%时，TraderAgent强制抛出HumanApprovalRequired异常")
    add_bullet(doc, "MockBroker隔离：所有交易只在本地MockBroker模拟，禁止连接真实券商API")

    # ── 十、第八步：回测引擎──────────────────────────────────────────
    add_title(doc, "十、第八步：回测引擎（解决「验证策略历史表现」的需求）", level=1)
    add_body(doc, "这一步的目标：用历史数据模拟交易，计算策略的Sharpe Ratio和最大回撤，"
             "客观评估策略是否可靠，而不是凭感觉。")
    add_body(doc, "回测流程：")
    add_numbered(doc, "输入：股票代码、回测时间段、初始资金")
    add_numbered(doc, "逐日遍历：每天触发六大Agent分析，得到buy/sell/hold信号")
    add_numbered(doc, "模拟交易：出现buy信号时，下一根K线开盘买入；出现sell信号时，下一根K线开盘卖出")
    add_numbered(doc, "扣除成本：滑点0.1%（模拟市场冲击）+ 手续费0.03%（双边共0.13%）")
    add_numbered(doc, "计算指标：Sharpe Ratio（收益风险比）、Max Drawdown（最大回撤）、年化收益率")
    add_body(doc, "代码位置：src/agent_platform/finance/backtesting.py（run_backtest函数）")
    add_body(doc, "E-01验收标准：要求Sharpe>0.5。实际结果：真实股票Sharpe=-0.440（未达标），"
             "合成数据Sharpe=+0.33。未达标原因：A股市场噪声大，纯技术指标信号在随机走势下"
             "难以稳定获得正Sharpe，这是金融领域的客观挑战。")

    # ── 十一、第九步：API和UI────────────────────────────────────────
    add_title(doc, "十一、第九步：API 和 UI（解决「让普通人能用」的需求）", level=1)
    add_body(doc, "这一步的目标：把整个系统封装成REST API，再做一个简单的网页前端，"
             "让不懂代码的用户也能使用这套系统。")
    add_body(doc, "FastAPI后端（src/agent_platform/api/main.py）：")
    add_bullet(doc, "11条REST路由：注册/登录、查行情、触发分析、查回测、提交订单、健康检查等")
    add_bullet(doc, "JWT认证：用户登录后返回Token，后续请求带Token访问")
    add_bullet(doc, "自动文档：访问 /docs 可查看交互式API文档（Swagger UI）")
    add_body(doc, "Streamlit前端（src/agent_platform/ui/app.py）：")
    add_bullet(doc, "输入股票代码 → 点击【开始分析】→ 展示六路Agent分析结果和综合信号")
    add_bullet(doc, "输入回测参数 → 点击【运行回测】→ 显示收益曲线和Sharpe/MaxDD指标")
    add_body(doc, "SQLite数据库：存储用户账号、分析记录、回测结果，轻量级单文件数据库。")

    # ── 十二、技术词典 ───────────────────────────────────────────────
    add_title(doc, "十二、技术词典——看不懂的术语查这里", level=1)
    add_kv_table(doc, [
        ("术语", "白话解释"),
        ("Agent",           "能自主完成任务的AI程序，能感知环境、做决策、调用工具、产生结果"),
        ("Harness",         "线束。这里指包裹Agent的生命周期框架，控制Agent的输入输出和运行过程"),
        ("ReAct",           "推理+行动。Agent的思考模式：先想（Reason）再做（Act）再看结果（Observe）"),
        ("DAG",             "有向无环图。描述多个任务之间依赖关系的图，箭头方向=执行顺序，无循环"),
        ("Guardrail",       "护栏。检查AI输出是否合规的规则，不合规就拦截"),
        ("CircuitBreaker",  "熔断器。连续失败N次后停止尝试，冷却一段时间后再恢复，防止雪崩"),
        ("Observability",   "可观测性。系统运行时能看到延迟/错误率/调用次数等指标"),
        ("TraceRecord",     "追踪记录。每次Agent运行的完整日志，包括输入输出/耗时/违规"),
        ("Kahn算法",        "拓扑排序算法，把图中节点按执行顺序排好，保证依赖关系正确"),
        ("ThreadPoolExecutor","线程池执行器。Python的并发工具，让多个任务同时执行（并行）"),
        ("JSONSchema",      "JSON格式的结构描述规范，用来验证JSON数据是否符合预定格式"),
        ("ReAct Loop",      "ReAct推理循环，Agent反复思考→行动→观察的过程，最多4次"),
        ("Sharpe Ratio",    "夏普比率。收益/风险比值，>1优秀，>0.5及格，<0不及格"),
        ("Max Drawdown",    "最大回撤。从历史最高点跌到最低点的最大跌幅，衡量最坏情况"),
        ("MACD",            "指数平滑异同移动平均线。判断趋势方向和强弱的技术指标"),
        ("RSI",             "相对强弱指数。判断超买（>70卖出信号）超卖（<30买入信号）"),
        ("MockBroker",      "模拟券商。本地假装是券商，执行买卖但不连接真实市场"),
        ("AkShare",         "A股数据接口Python库，免费获取股票行情/财务数据"),
        ("FastAPI",         "高性能Python Web框架，用来构建REST API，自带Swagger文档"),
        ("Streamlit",       "Python数据可视化Web框架，几行代码做出交互式网页"),
        ("JWT",             "JSON Web Token。用户登录后的身份令牌，用于API鉴权"),
        ("bcrypt",          "密码哈希算法，存储用户密码时不存原文，存哈希值"),
    ])

    # ── 十三、如何运行项目 ─────────────────────────────────────────────
    add_title(doc, "十三、如何运行这个项目？", level=1)
    add_body(doc, "前提：已安装Python 3.11+，在项目根目录下执行以下命令：")
    add_numbered(doc, "安装依赖：.venv/Scripts/python.exe -m pip install -r requirements.txt")
    add_numbered(doc, "配置密钥：把DEEPSEEK_API_KEY和ANTHROPIC_API_KEY写入 .env 文件")
    add_numbered(doc, "运行全部测试：.venv/Scripts/pytest tests/ （预期358个测试全通过）")
    add_numbered(doc, "启动API服务：.venv/Scripts/uvicorn src.agent_platform.api.main:app --port 8002")
    add_numbered(doc, "启动UI前端：.venv/Scripts/streamlit run src/agent_platform/ui/app.py")
    add_numbered(doc, "生成文档：.venv/Scripts/python.exe Scripts/generate_project_documents.py")

    # ── 十四、向导师汇报时这样说 ─────────────────────────────────────
    add_title(doc, "十四、向导师汇报时，这样介绍这个项目", level=1)
    add_body(doc, "这个项目解决的核心问题是：如何让大模型在企业级场景中可靠、可控、可审计地运行。")
    add_body(doc, "我采用了2026年最新的Harness Engineering V2.0范式，核心思想是「Agent = Model + Harness」——"
             "把大模型包裹在一个生命周期框架里，所有输入输出都经过护栏检查，"
             "3次连续失败自动熔断，每次运行都有完整的追踪日志。")
    add_body(doc, "在这个基础平台上，我搭建了证券金融分析应用：6个专业Agent分工分析"
             "技术面、基本面、行业面和市场风格，SynthesisAgent汇总评分输出买卖信号，"
             "再通过回测引擎验证历史表现。整个系统用FastAPI暴露11条API路由，"
             "Streamlit提供可视化界面，358个测试全部通过，幻觉阻断率100%。")
    add_body(doc, "主要技术挑战：E-01验收标准（Sharpe>0.5）未达标，"
             "真实股票回测Sharpe=-0.440。原因是A股市场噪声大，纯技术指标策略难以稳定盈利。"
             "我没有通过选择有利时段来人为美化结果，而是诚实记录并提出改进方向："
             "引入多因子联合信号可望改善这一指标。")

    doc.save(BEGINNER)
    print(f"[OK] 已生成：{BEGINNER}")


# ─────────────────────────────────────────────────────────────────────────────
# Document 3: 工作总结
# ─────────────────────────────────────────────────────────────────────────────
def build_work_summary() -> None:
    doc = Document()
    configure_document(doc)

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("工作总结报告")
    set_run_font(run, size=18, bold=True)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_para.add_run(
        "项目：构建通用Agent平台及证券金融分析应用  "
        "实习生：米志轩  时间：2026年7月29日—8月5日")
    set_run_font(run2, size=11)
    doc.add_paragraph()

    # 12行 × 4列 Table Grid 表格
    ROWS = [
        ("序号", "工作内容",                         "完成情况", "收获与备注"),
        ("01", "阅读任务书，理解Harness Engineering V2.0范式，搭建项目脚手架（目录结构/虚拟环境/依赖）",
               "完成", "理解Agent=Model+Harness核心理念，建立工程框架"),
        ("02", "实现AgentHarness核心类：Pre-flight→Loop→Post-flight生命周期，CircuitBreaker熔断器",
               "完成", "掌握熔断器模式（max_failures=3，cooldown=300s）"),
        ("03", "实现五大Guardrail护栏体系：JSONSchema/SourceAttribution/RateLimiter/KeywordBlocker/CrossValidator",
               "完成", "幻觉阻断率100%，误报率0%，实验数据验证"),
        ("04", "实现AgentRuntime ReAct循环引擎：Reason→Act→Observe，max_steps=4，ToolRegistry",
               "完成", "ReAct范式是Agent具备自主推理能力的核心"),
        ("05", "迁移主工作流至 LangGraph StateGraph：并行扇出、interrupt/resume 人工审批、SQLite checkpoint",
               "完成", "LangGraph 已成为证券分析主编排引擎"),
        ("06", "实现数据层和技术指标：AkShare行情接入、10类指标函数（MA/EMA/MACD/RSI/BOLL/KDJ等）",
               "完成", "熟悉常用技术分析指标的计算逻辑"),
        ("07", "实现六大专业Agent：Technical/Fundamental/Industry/MarketRegime/Synthesis/Trader",
               "完成", "分工协作模式显著提升分析质量和可维护性"),
        ("08", "实现回测引擎：滑点0.1%/手续费0.03%，Sharpe/MaxDrawdown计算；E-01验收测试",
               "部分完成", "真实股票Sharpe=-0.440，合成数据+0.33；E-01未达标，分析原因"),
        ("09", "实现FastAPI后端（11条路由）和Streamlit前端原型，SQLite持久化，JWT认证",
               "完成", "全链路可用：注册/登录/分析/回测/订单/可观测"),
        ("10", "补全全套测试用例，运行358个测试，全部通过（0 failed，0 error）",
               "完成", "测试驱动保障代码质量，回归测试用时约2分钟"),
        ("11", "整理项目文档（项目总结报告、小白说明书、工作总结），回顾并记录项目全貌",
               "完成", "文档化是工程能力的重要体现，也是知识沉淀的关键"),
    ]

    table = doc.add_table(rows=len(ROWS), cols=4)
    table.style = "Table Grid"

    col_widths = [Cm(1.2), Cm(9.0), Cm(2.8), Cm(5.0)]
    for i, row_data in enumerate(ROWS):
        row = table.rows[i]
        is_header = (i == 0)
        fill = "2E74B5" if is_header else ("D9E1F2" if i % 2 == 1 else "FFFFFF")
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            set_cell_margins(cell)
            if is_header:
                set_cell_shading(cell, "2E74B5")
                set_cell_text(cell, cell_text, bold=True, size=10.5,
                              align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                set_cell_shading(cell, fill)
                align = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
                set_cell_text(cell, cell_text, bold=False, size=10, align=align)

    # 设置列宽
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width

    doc.add_paragraph()
    add_body(doc, "工作总结：本次实习历时8天，从零搭建了完整的通用Agent平台及证券金融分析应用。"
             "项目涵盖Harness Engineering V2.0范式、ReAct推理循环、DAG多Agent并行调度、"
             "五大护栏安全体系、六大专业Agent、回测引擎和REST API，"
             "共完成358个测试用例，全部通过。"
             "E-01（Sharpe>0.5）未达标，已分析原因并提出后续改进方向。")

    doc.save(WORK_SUMMARY)
    print(f"[OK] 已生成：{WORK_SUMMARY}")


# ─────────────────────────────────────────────────────────────────────────────
# Validate & Main
# ─────────────────────────────────────────────────────────────────────────────
def validate_documents() -> bool:
    ok = True
    for path in (SUMMARY, BEGINNER, WORK_SUMMARY):
        if path.exists() and path.stat().st_size > 5000:
            print(f"[PASS] {path.name}  ({path.stat().st_size:,} bytes)")
        else:
            print(f"[FAIL] {path.name}  ← 文件不存在或体积过小")
            ok = False
    return ok


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    print("=" * 60)
    print("生成项目文档 V2 ...")
    print("=" * 60)
    build_project_summary()
    build_beginner_guide()
    build_work_summary()
    print()
    print("─" * 60)
    print("验证文档 ...")
    ok = validate_documents()
    print("─" * 60)
    if ok:
        print("[DONE] 三份文档全部生成成功！")
    else:
        print("[ERROR] 部分文档生成失败，请检查以上错误。")
        raise SystemExit(1)


if __name__ == "__main__":
    main()

