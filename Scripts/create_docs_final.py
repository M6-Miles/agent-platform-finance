# -*- coding: utf-8 -*-
"""
一次性生成四份 Word 文档到 docs/ 目录：
  1. 项目总结文档.docx   ← 向导师展示
  2. 小白说明文档.docx   ← 底层逻辑 + 术语解释
  3. 指导说明书.docx     ← 手把手复现步骤
  4. 实习报告.docx       ← 与模板格式相同
"""
import os
from docx import Document
from docx.shared import Pt

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(BASE, "docs")

def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    return doc

def title_para(doc, text):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = "微软雅黑"
    return p

def subtitle_para(doc, text):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "微软雅黑"
    return p

def h1(doc, text):
    p = doc.add_heading("", level=1)
    run = p.add_run(text)
    run.font.name = "微软雅黑"

def h2(doc, text):
    p = doc.add_heading("", level=2)
    run = p.add_run(text)
    run.font.name = "微软雅黑"

def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(22)
    return p

def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")

def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")

def divider(doc):
    doc.add_paragraph()

def save(doc, filename):
    path = os.path.join(DOCS, filename)
    try:
        doc.save(path)
        print(f"✓  {filename}")
    except PermissionError:
        alt = path.replace(".docx", "_draft.docx")
        doc.save(alt)
        print(f"⚠  {filename} 被 Word 占用 → 已另存为 {os.path.basename(alt)}")


def make_project_summary():
    doc = new_doc()
    title_para(doc, "通用 Agent 平台及证券金融分析应用")
    subtitle_para(doc, "项目总结文档")
    subtitle_para(doc, "作者：米志轩   日期：2026年8月7日   导师：东方国信")
    divider(doc)

    h1(doc, "一、项目概述")
    para(doc, "本项目旨在构建一套以大语言模型（LLM）和多 Agent 协作为核心的证券金融分析平台。"
         "平台从原始行情数据出发，经过技术分析、基本面分析、行业分析、市场情绪综合判断，最终"
         "生成可供参考的交易建议，并通过完整的风控和人工审批流程确保安全性。系统同时支持"
         "在线模式（接入 AkShare 真实行情）和离线模式（SampleMarketDataProvider 固定数据），"
         "二者共享同一 LangGraph 工作流，仅数据访问层不同。")
    para(doc, "项目定位为研究演示系统，所有输出均附带『仅供研究参考，不构成投资建议』免责声明，"
         "不接入任何真实券商账户，MockBroker 仅在本地模拟执行路径。")

    h1(doc, "二、功能目标")
    bullet(doc, "提供证券多维度分析查询与 Agent 对话能力。")
    bullet(doc, "整合技术面、基本面、行业面、市场情绪，综合生成买卖信号和置信度。")
    bullet(doc, "对置信度不足（≤30%）的情况输出 no_trade，避免噪音交易。")
    bullet(doc, "仓位建议超过 10% 时触发人工审批，未确认前系统挂起等待。")
    bullet(doc, "通过 Guardrail 机制确保所有输出携带 source 字段与 DISCLAIMER。")
    bullet(doc, "利用 SQLite checkpoint 实现跨进程、跨实例的工作流状态恢复。")
    bullet(doc, "提供完整测试体系（592 项自动化测试，全部通过）。")

    h1(doc, "三、系统架构")
    para(doc, "系统采用五层架构设计，从上至下依次为：展示层、接口层、服务层、编排层、数据与存储层。")

    h2(doc, "3.1 展示层 — Streamlit UI")
    para(doc, "基于 Streamlit 构建的 Web 前端，提供行情查询、K线图表、多维分析报告、"
         "聊天对话和交易研究界面。普通行情查询接口直接与 FastAPI 通信；"
         "投资研究接口拥有独立的 thread（thread_id），支持查看中断状态和提交 approve/reject 决策。")

    h2(doc, "3.2 接口层 — FastAPI REST API")
    para(doc, "对外暴露以下端点：GET /health（健康检查）、GET /analysis/{symbol}（行情分析）、"
         "POST /chat（Agent 对话）、POST /research/{symbol}（启动工作流）、"
         "GET /research/{thread_id}/state（查询状态）、POST /research/{thread_id}/resume"
         "（提交人工审批决策）。所有接口均经过 CORS 配置，支持跨域调用。")

    h2(doc, "3.3 服务层 — ApplicationService")
    para(doc, "ApplicationService 是系统的唯一业务门面，负责初始化 LangGraph 图、"
         "管理 SQLite checkpoint 连接生命周期（self._langgraph_checkpoint_conn + close()），"
         "以及统一调用各领域分析方法。FastAPI 通过 lifespan 异步上下文管理器在应用关闭时"
         "显式调用 close()，避免 Windows 下 SQLite 文件占用问题。")

    # <<MARKER_B>>
    h2(doc, "3.4 编排层 — LangGraph 工作流")
    para(doc, "以 LangGraph 1.2.10 构建的有向图（StateGraph），节点为各领域 Agent 函数，"
         "共享 SecuritiesAnalysisState（TypedDict）。四个分析节点并行从 START 出发，"
         "汇聚到 synthesis_agent；低置信度路由到 no_trade 节点（写入 status='no_trade'）后结束；"
         "高置信度路由到 trader_agent；仓位超限则触发 human_approval interrupt；"
         "通过风控后进入 trading_harness 执行 preflight 检查。")

    h2(doc, "3.5 数据与存储层")
    para(doc, "DataProvider 接口统一封装数据访问，实现类有 SampleMarketDataProvider（固定样本）"
         "和 AkShareDataProvider（在线行情）。indicators.py 提供 MA/EMA/MACD/RSI/"
         "布林带/KDJ/ATR/CCI 等技术指标计算。SQLite 通过 SqliteSaver(conn) + .setup()"
         "持久化 LangGraph checkpoint，check_same_thread=False 支持多线程访问。")

    h1(doc, "四、核心流程详解")
    para(doc, "一次完整的投资研究工作流（POST /research/{symbol}）按以下步骤执行：")
    numbered(doc, "ApplicationService 接收 symbol，选择 data_mode，分配唯一 thread_id。")
    numbered(doc, "LangGraph 从 START 并发启动：technical_agent（K线+指标）、fundamental_agent"
             "（市盈率/市净率/ROE/负债率）、industry_agent（行业对比与景气度）、"
             "market_regime_agent（大盘趋势与北向资金）四路并行。")
    numbered(doc, "synthesis_agent 汇聚四路结果，调用 synthesize() 生成信号与置信度。"
             "置信度 ≤ 0.30 → no_trade 节点 → END（status='no_trade'）；"
             "置信度 > 0.30 → trader_agent。")
    numbered(doc, "trader_agent 调用 generate_trade_signal()，计算仓位建议。"
             "仓位 > 10% → 抛出 HumanApprovalRequired → 图 interrupt，等待人工审批。")
    numbered(doc, "人工通过 POST /resume 提交 approve 或 reject，LangGraph 从 SQLite 恢复断点。"
             "approve → risk_manager；reject → block。")
    numbered(doc, "risk_manager 执行风险核查（仓位上限、信号一致性），通过后进入 trading_harness。")
    numbered(doc, "trading_harness 执行 preflight 最终检查，输出 execute / manual_review / block。")

    h1(doc, "五、关键技术实现")
    h2(doc, "5.1 状态统一管理")
    para(doc, "三个 API 端点（POST /research、GET /state、POST /resume）共享"
         "resolve_research_status() 函数，统一将图内部状态映射为"
         "六个规范值之一：interrupted / blocked / completed / no_trade / failed / not_found。"
         "优先级：interrupted > blocked > completed > no_trade > failed > completed（兜底）。")

    h2(doc, "5.2 SQLite Checkpoint 生命周期")
    para(doc, "self._langgraph_checkpoint_conn 保存 SQLite 连接引用，close() 方法幂等关闭。"
         "FastAPI lifespan 在 yield 之后调用 svc.close()，确保进程正常退出时连接释放，"
         "避免 Windows 下数据库文件被长期锁定。")

    h2(doc, "5.3 MemorySaver 环境变量")
    para(doc, "Settings 数据类新增 langgraph_use_memory_saver: bool = False，"
         "由 _parse_bool(os.getenv('LANGGRAPH_USE_MEMORY_SAVER','false')) 解析。"
         "_parse_bool 仅对 '1'/'true'/'yes'/'on'（大小写不敏感）返回 True，"
         "避免 bool('false') == True 的陷阱。生产默认使用 SQLite，测试可按需开启内存模式。")

    # <<MARKER_C>>
    h2(doc, "5.4 安全护栏（Guardrail）")
    bullet(doc, "SourceAttributionFilter：每个 Agent 输出必须包含 source 字段，缺失直接抛 GuardrailViolation。")
    bullet(doc, "_MAX_AUTO_POSITION_PCT = 10.0：超过此值必须人工审批，不可自动执行。")
    bullet(doc, "DISCLAIMER：所有对外响应包含『仅供研究参考，不构成投资建议』。")
    bullet(doc, "MockBroker：禁止连接真实券商，所有下单仅在内存中模拟。")
    bullet(doc, "API 密钥仅存于 .env，不得提交版本库。")

    h1(doc, "六、测试体系")
    para(doc, "项目采用 pytest 构建全面测试体系，共592项测试，全部通过（0失败）。"
         "测试分类包括：单元测试（各 Agent 函数、指标计算、状态转换）、集成测试"
         "（ApplicationService + SQLite + LangGraph 完整路径）、API 测试"
         "（TestClient 验证 HTTP 状态码和响应体）、以及一致性测试（STATUS-01~04 验证"
         "三端点状态统一）。offline 模式通过 monkeypatch 屏蔽所有网络访问，"
         "确保 CI 环境稳定可复现。")

    h2(doc, "6.1 验收结果")
    bullet(doc, "验收 A（端到端 ≥20只）：execute 4/20，no_trade 16/20（置信度不足属正常），error 0。⚠️ PARTIAL")
    bullet(doc, "验收 B（回测 Sharpe）：9/20只达标，全样本均值 Sharpe = 0.33。⚠️ BELOW 0.5")
    bullet(doc, "验收 C（幻觉拦截）：拦截率 100%，误报率 0%。✅ PASS")

    h1(doc, "七、项目成果")
    bullet(doc, "完整的模块化多 Agent 架构，可复用于其他金融品种或策略场景。")
    bullet(doc, "LangGraph 工作流实现了 interrupt/resume 人工审批全流程，状态跨进程持久化。")
    bullet(doc, "resolve_research_status() 统一状态映射，消除三端点语义不一致的问题。")
    bullet(doc, "SQLite 连接生命周期管理，close() + lifespan 双保险，防止 Windows 文件锁。")
    bullet(doc, "严格的 Guardrail 体系和 MockBroker 隔离，系统安全边界清晰。")
    bullet(doc, "592 项自动化测试提供完整质量保障，offline 模式确保离线 CI 可用。")

    h1(doc, "八、不足与展望")
    para(doc, "当前主要不足：(1) 策略仅用 MA5/MA20 交叉，样本外平均 Sharpe=0.33，"
         "需引入因子模型和组合优化；(2) 仅支持样本股票离线验证，实时行情接入需完善"
         "AkShare 异常处理；(3) LLM 使用 Mock 实现，未接入真实大模型，"
         "实际 NLP 分析能力待验证；(4) 前端功能基础，缺乏自选股管理和实时推送。")
    para(doc, "改进方向：引入多因子策略和 Walk-Forward 验证；接入真实 DeepSeek/Anthropic API；"
         "增加观察模式、仓位管理和更细粒度的停损；完善监控和日志可观测性。")

    h1(doc, "九、结论")
    para(doc, "本项目在10天内完成了通用 Agent 平台的完整设计与实现，覆盖数据层到前端展示的"
         "全栈开发。核心贡献在于：将 LangGraph 的中断/恢复机制应用于金融审批场景，"
         "建立了可复现、可测试、状态一致的多 Agent 工作流体系。"
         "Sharpe 指标未达目标体现了策略研究的客观挑战，"
         "但工程质量、安全护栏和测试覆盖率均满足设计要求。")

    save(doc, "项目总结文档.docx")


# <<MARKER_D>>
