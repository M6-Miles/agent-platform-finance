# -*- coding: utf-8 -*-
"""Generate the two current Word deliverables from verified project facts."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FONT = "Microsoft YaHei"
BLUE = "1F4E79"
LIGHT = "EAF2F8"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
RED = "FCE4D6"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), fill)
    tcPr.append(el)


def borders(cell, color="B7C9D6"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        x = OxmlElement(f"w:{side}")
        x.set(qn("w:val"), "single")
        x.set(qn("w:sz"), "4")
        x.set(qn("w:color"), color)
        b.append(x)
    tcPr.append(b)


def set_font(run, size=10.5, bold=False, color="000000"):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rf.set(qn(key), FONT)


def setup(doc, title):
    sec = doc.sections[0]
    sec.top_margin = Cm(1.7); sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.9); sec.right_margin = Cm(1.9)
    normal = doc.styles["Normal"]
    normal.font.name = FONT; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color in (("Title", 22, BLUE), ("Heading 1", 15, BLUE), ("Heading 2", 12.5, "2F5597"), ("Heading 3", 11, "44546A")):
        st = doc.styles[name]
        st.font.name = FONT; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); set_font(r, 22, True, BLUE)
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("企业级研究型 Agent 平台项目交付材料 | 版本：2026-08-14"); set_font(r, 10, False, "667085")
    doc.add_paragraph()


def p(doc, text, bold_prefix=None):
    para = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = para.add_run(bold_prefix); set_font(r, bold=True)
        r = para.add_run(text[len(bold_prefix):]); set_font(r)
    else:
        r = para.add_run(text); set_font(r)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        r = para.add_run(item); set_font(r)


def steps(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        r = para.add_run(item); set_font(r)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; shade(c, BLUE); borders(c)
        r = c.paragraphs[0].add_run(h); set_font(r, 9.5, True, "FFFFFF")
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""; borders(cells[i]); shade(cells[i], "FFFFFF")
            r = cells[i].paragraphs[0].add_run(str(value)); set_font(r, 9)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in t.rows:
            for cell, width in zip(row.cells, widths): cell.width = Cm(width)
    doc.add_paragraph()
    return t


def callout(doc, label, text, fill=LIGHT):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); shade(c, fill); borders(c, "9FBAD0")
    c.text = ""; r = c.paragraphs[0].add_run(label + "："); set_font(r, 10.5, True, BLUE)
    r = c.paragraphs[0].add_run(text); set_font(r, 10.5)
    doc.add_paragraph()


def code(doc, text):
    t = doc.add_table(rows=1, cols=1); c = t.cell(0, 0); shade(c, "F3F4F6"); borders(c, "D0D5DD"); c.text = ""
    r = c.paragraphs[0].add_run(text); set_font(r, 9, False, "344054")
    doc.add_paragraph()


def summary_doc():
    d = Document(); setup(d, "项目总结文档")
    callout(d, "执行摘要", "本项目实现了一个以 Harness、Loop、Graph 为核心范式的通用 Agent 平台，并落地证券金融分析、天气分析、模拟盘和真实 LLM 回放实验。当前交付物是可运行的研究级工程原型，不是连接真实券商的生产交易系统。")
    d.add_heading("一、项目目标与边界", 1)
    p(d, "项目目标是把“模型调用”升级为可审计、可恢复、可测试的 Agent 应用：系统能够接收用户任务，调用工具取得数据，由多个专业 Agent 分析，再经过结构化校验、来源校验、风险控制和人工审批，最终返回研究结果。")
    p(d, "边界必须明确：系统只生成研究建议和模拟交易结果；MockBroker 不连接券商、不提交真实订单；真实行情和天气属于外部网络依赖；回测指标如 Sharpe 未达标时按事实记录，不通过修改公式或样本伪造达标。")
    d.add_heading("二、总体架构", 1)
    table(d, ["层次", "主要组件", "职责"], [
        ("展示层", "frontend_prototype.html；Streamlit 备用界面", "登录、行情、证券分析、深度投研、天气、模拟盘和状态展示"),
        ("接口层", "FastAPI", "HTTP 路由、鉴权、参数校验、统一错误响应、研究任务和状态查询"),
        ("编排层", "LangGraph 1.x；SQLite Checkpoint", "节点、条件路由、中断、恢复、审批和任务状态持久化"),
        ("Agent 层", "Technical/Fundamental/Industry/Regime/Bull-Bear/Synthesis/Trader/Risk", "分工分析、综合判断、交易建议和风险拦截"),
        ("数据层", "Provider Factory；AkShare/腾讯公开行情/Open-Meteo/样例", "按模式读取数据，记录 source、data_status、updated_at，并支持受控降级"),
        ("存储层", "SQLite store、用户表、会话、checkpoint、模拟盘记录", "身份、历史、任务恢复和审计数据"),
    ], [2.3, 5.2, 8.5])
    d.add_heading("三、核心设计范式", 1)
    d.add_heading("3.1 Harness：模型外部的质量控制线", 2)
    p(d, "Harness 不是模型本身，而是包围模型调用的工程控制层。它统一处理输入、Provider、输出解析、Schema 校验、Guardrail、重试、日志和指标。这样模型输出不再直接进入交易建议流程。")
    d.add_heading("3.2 Loop：让 Agent 能持续完成任务", 2)
    p(d, "Loop 由规划、工具调用、观察结果、反思判断和继续或结束五个要素组成。每轮循环都有状态，遇到工具失败可以重试，遇到风险或需要人决定时可以停止。")
    d.add_heading("3.3 Graph：把流程变成可审计的状态机", 2)
    p(d, "证券深度投研使用 LangGraph 节点和条件边表达流程。图中节点负责单一动作，条件边根据 status、置信度和风险结果选择下一步；interrupt 将流程暂停并等待 approve/reject；SQLite checkpoint 使服务重启后仍能恢复。")
    d.add_heading("四、证券分析工作流", 1)
    steps(d, [
        "接收 symbol 和 data_mode，校验输入并建立 thread_id。",
        "四个 Specialist Agent 并行或按编排策略读取行情、基本面、行业和市场状态。",
        "Bull/Bear 对关键结论进行多空观点对照，Synthesis 生成结构化综合报告。",
        "Trader 生成参考价、止损价、交易信号和建议仓位；Risk Manager 按止损距离计算单笔账户风险并限制在 2% 内。",
        "需要人工决策时触发 human approval；批准后执行模拟动作，拒绝则进入 block；置信度不足进入 no_trade。",
        "所有结果写入状态和 checkpoint，API 返回 completed、interrupted、no_trade、failed 等真实状态。",
    ])
    d.add_heading("五、数据真实性与降级策略", 1)
    table(d, ["模式", "数据来源", "允许的表述", "限制"], [
        ("offline", "SampleMarketDataProvider / offline_sample_data.py", "内置样例数据", "不代表实时行情"),
        ("auto/联网", "Provider Factory 选择 AkShare、腾讯公开接口等", "联网行情/公开接口返回", "依赖网络、接口可用性和交易时段"),
        ("fallback", "主数据源失败后受控切换", "降级数据，并记录 fallback_reason", "不能把降级结果标为实时"),
        ("unavailable", "无法取得或校验失败", "数据不可用", "应停止相关分析，不静默填零"),
    ], [2.3, 5.1, 4.2, 4.4])
    p(d, "每个关键结果应包含 source、data_status、updated_at；价格、昨收、交易时间和交易日范围进行一致性校验。真实数据和样例数据在前端徽标、报告和 API 中明确区分。")
    d.add_heading("六、模拟盘与安全边界", 1)
    p(d, "模拟盘使用 MockBroker 和本地 SQLite 记录现金、持仓、订单、成交、净值和异常。它用于验证订单状态机、幂等、恢复和前端交互，不等价于真实交易。2026-08-14 已记录首个有效真实交易日：000001 和 600519 均为腾讯公开 live 行情且无降级，当前为 1/7；其余天数必须由自然时间积累。")
    d.add_heading("七、真实 DeepSeek Harness 回放证据", 1)
    callout(d, "证据边界", "真实回放属于固定离线评测，不等于生产用户流量。Mock 固定集的 100% 拦截结果不能替代真实模型指标。", AMBER)
    d.add_heading("7.1 100 条真实结果", 2)
    p(d, "2026-08-14 使用固定人工标注评测集 enterprise_harness_100_v1 执行了 100 条真实 DeepSeek 回放。每条任务只调用模型一次，再对同一响应执行 OFF/ON 对照。")
    table(d, ["指标", "实测值", "解释"], [
        ("样本数", "100", "40 正常、20 缺字段、15 违规、10 注入、10 数据风险、5 格式异常"),
        ("Schema 合格率", "67.0%", "100 条中 67 条通过 OFF 结构检查"),
        ("标签匹配率", "91.0%", "实际结果与人工预期类别一致的比例"),
        ("违规拦截召回率", "100%", "预期违规样本全部被识别并拦截"),
        ("违规拦截精确率", "93.75%", "实际违规拦截中属于预期违规的比例"),
        ("正常请求误报率", "0%", "40 条正常请求没有被错误拦截"),
        ("人工审核率", "11.0%", "输入风险检测命中的比例"),
        ("固定事实错误", "0/40", "40 条携带事实快照的正常任务未观测到事实篡改"),
        ("无效下游动作资格", "OFF 60 -> ON 0", "Harness 阻止结构或合规不合格结果进入下游"),
        ("P50/P95/P99", "2.144 / 2.821 / 5.002 秒", "真实调用延迟分布"),
    ], [4.0, 4.0, 8.0])
    callout(d, "真实结论", "真实模型的 100 条结果显示，本轮违规拦截召回率为 100%、正常误报率为 0%，但仍有 33% 的输出未通过结构检查。40 条事实核验样本未观测到事实错误，因此事实错误阻断率没有有效分母，应记为 N/A，不能声称已证明真实幻觉率为零。当前只完成 1 个自然日，7 天证据状态仍为 collecting。", RED)
    p(d, "新版评测集给正常任务附带固定事实快照，并保存脱敏后的结构化输出逐字段核验；只有完成核验的样本才进入幻觉率分母。同时统计 Harness OFF/ON 下输出是否有资格触发下游动作，以计算被阻止的无效调用。该统计不执行真实交易或业务 API。旧报告缺少这些字段，不能倒推新指标。")
    p(d, "结构化输出兼容性补充：deepseek-chat 在本环境对 response_format 请求出现 HTTP 400，因此系统不会强行启用 JSON Mode；deepseek-v4-flash 的 3 条冒烟测试可正常返回 JSON，但 P95 延迟约 11.08 秒。系统现在按模型能力自动选择，优先保证兼容性，不把协议错误伪装成模型质量指标。")
    d.add_heading("八、回测与未达标项", 1)
    p(d, "回测保留原 Sharpe 公式、年化方式、无风险利率、0.5 阈值和 MA baseline。当前已完成成本、滑点、印花税、连续仓位、Walk-forward、样本外对比和基准收益记录，但策略研究结果仍未稳定达到 Sharpe 0.5。正式基线多因子样本外均值约为 -0.337；本轮预先定义的稳健选参挑战方案实测为 -0.499，未带来改善，因此没有替换正式基线。")
    d.add_heading("九、质量验证", 1)
    table(d, ["门禁", "结果/状态", "证据"], [
        ("单元与集成测试", "1621 项收集；1620 passed，1 skipped，0 failed", "2026-08-14 pytest"),
        ("静态检查", "pyflakes、compileall 已通过", "命令输出"),
        ("前端语法", "已通过", "Scripts/check_frontend_syntax.js"),
        ("敏感信息", "API Key 不写入代码、文档、日志或 Git", ".env 与报告审计"),
        ("真实券商", "未连接", "MockBroker 与 AST/集成测试"),
    ], [4.0, 5.0, 7.0])
    d.add_heading("十、导师验收建议", 1)
    steps(d, [
        "复制 .env.example 为 .env，保持 LLM_PROVIDER=mock，先验证离线流程。",
        "启动 FastAPI，打开 http://127.0.0.1:8003/docs 查看接口。",
        "使用 DEMO001 + offline 验证四个 Agent、source 和 data_status。",
        "使用 auto 模式验证联网成功、降级和 unavailable 三种状态。",
        "验证深度投研 interrupted 后 approve/reject、服务重启恢复和 no_trade。",
        "运行 pytest，并查看 docs/experiments 中脱敏后的真实回放报告。",
        "检查模拟盘仅产生本地 MockBroker 记录，没有真实订单。",
    ])
    d.add_heading("十一、后续计划", 1)
    bullets(d, [
        "扩大真实 LLM 回放任务和多日期运行，形成可统计的 Schema、拦截、延迟和重试样本。",
        "自然运行模拟盘至少 7-14 个真实交易日后，再决定是否满足时间证据要求。",
        "继续采用严格样本外和多基准验证改进策略，保留原始 baseline 和公式。",
        "生产化前补充密钥管理、用户级隔离、监控告警、备份恢复、审计保留策略和部署演练。",
    ])
    d.add_heading("十二、免责声明", 1)
    p(d, "本系统用于 Agent 工程、数据流程、风控和研究方法演示。证券分析结果不构成投资建议；联网数据可能延迟、缺失或受第三方接口限制；严禁把模拟盘结果当作真实收益，也不得在未完成生产安全审查前接入真实交易。")
    d.save(DOCS / "项目总结文档.docx")


def beginner_doc():
    d = Document(); setup(d, "项目小白说明文档")
    callout(d, "当前事实", "截至 2026-08-14：1621 项测试收集，1620 passed、1 skipped、0 failed；正式样本外 Sharpe 约 -0.337，未达到 0.5；真实行情模拟盘证据为 1/7 天。", AMBER)
    callout(d, "先看结论", "你可以把这个项目理解成“一个会调用数据、多个专业分析员协作、还会自我检查并等待人批准的研究系统”。它能做研究演示和模拟交易，但不会自动拿你的钱去真实买股票。", GREEN)
    d.add_heading("第一部分：先建立整体画面", 1)
    d.add_heading("1. 这个项目到底做什么？", 2)
    p(d, "用户在网页中输入股票代码，系统取得行情和研究数据，然后分别从技术、基本面、行业和大盘环境四个角度分析，再综合成报告，给出研究性质的交易信号。若信号涉及模拟交易，风控模块先检查，必要时停下来等人批准。天气 Demo 复用了同样的“输入—调用工具—返回结构化结果—标注来源”思想。")
    d.add_heading("2. 前端和后端是什么？", 2)
    p(d, "前端就是你看见和点击的网页：输入框、按钮、图表、状态卡片都属于前端。项目中的主要前端是 frontend_prototype.html，备用前端是 Streamlit 页面。后端是藏在电脑后台的 Python 程序：它接收前端请求，计算、调用数据、保存记录，再把结果返回。")
    p(d, "一次点击的完整链路是：浏览器按钮 → JavaScript 发出 HTTP 请求 → FastAPI 路由接收 → Service 调用 Agent 和 Provider → SQLite 保存状态 → JSON 返回 → 浏览器更新卡片。HTTP 是浏览器和后端传话的规则；JSON 是双方都能读懂的结构化文字。")
    d.add_heading("3. 本项目中的三种运行方式", 2)
    table(d, ["方式", "你看到什么", "适合什么时候"], [
        ("offline", "固定、可重复的内置样例", "没有网络，学习和测试"),
        ("auto/联网", "公开接口返回的真实或延迟数据", "验证数据接入，必须接受网络失败"),
        ("Mock LLM", "固定的模型替身输出", "默认安全演示，不产生模型费用"),
    ], [3.2, 6.0, 5.8])
    d.add_heading("第二部分：从零解释专业名词", 1)
    terms = [
        ("Python", "一种编程语言", "后端、Agent、数据处理和测试都用它", "像用中文写步骤，但计算机能执行"),
        ("虚拟环境 .venv", "项目专用的小型 Python 环境", "避免不同项目的依赖互相打架", "项目 A 用旧版本，项目 B 用新版本，彼此不冲突"),
        ("依赖", "项目需要安装的第三方软件包", "FastAPI、pandas、LangGraph 等不是 Python 自带的", "像做菜前要准备锅、刀和食材"),
        ("API", "程序之间约定好的入口", "前端用 API 请求分析、行情和天气", "像银行柜台的业务窗口，不需要知道柜台内部怎么做"),
        ("FastAPI", "用 Python 编写 HTTP API 的框架", "定义 /health、/quote、/research 等后端入口", "接待请求、校验参数、返回 JSON"),
        ("JSON", "由键和值组成的数据格式", "Agent 结果、API 响应和报告都使用它", "{ price: 10.2, source: 'sample' }"),
        ("SQLite", "一个存放在本地文件里的数据库", "保存用户、会话、订单、checkpoint", "不需要单独安装数据库服务器的 Excel 替代品，但更适合程序"),
        ("Agent", "能按照目标使用工具完成任务的程序角色", "四个专业 Agent 分别负责不同分析", "像分工明确的研究员，不是一个万能函数"),
        ("LLM", "Large Language Model，大语言模型", "负责理解问题、生成文字或结构化结果", "DeepSeek 是一种 LLM；Mock 是它的本地替身"),
        ("Provider", "数据或模型的供应商适配器", "把不同来源统一成项目能理解的接口", "换数据源时不必重写所有 Agent"),
        ("Factory", "根据配置创建正确对象的工厂函数", "auto 根据环境变量选择 AkShare 等 Provider", "像前台根据票种发给你不同窗口"),
        ("Schema", "规定 JSON 必须有哪些字段和类型", "防止模型漏字段或返回乱七八糟的文字", "表格模板：姓名必须是文字，年龄必须是数字"),
        ("Guardrail", "阻止危险或不合规输出的护栏", "检查来源、违禁词、风险和结构", "像高速公路护栏，不能保证车永远正确，但能阻止冲出路面"),
        ("Harness", "包住模型调用的一整套测试和控制外壳", "统一做重试、解析、脱敏、指标和 Guardrail", "像模型进入生产前的质检线"),
        ("Loop", "任务反复执行直到完成或停止的循环", "规划→工具→观察→反思→继续/结束", "像研究员查资料、发现缺数据、再查一次"),
        ("Graph/DAG", "用节点和连线表示的流程图；DAG 不允许绕回形成环", "表示投研步骤和条件分支", "先做数据分析，再综合，再风控"),
        ("LangGraph", "专门管理 Agent 状态图的框架", "节点、条件边、暂停和恢复都由它管理", "让复杂流程可视化、可恢复、可测试"),
        ("Node", "图中的一个动作步骤", "技术分析、综合、风控各是节点", "一个函数只完成一个清晰动作"),
        ("Checkpoint", "某个时刻保存的任务快照", "服务重启后恢复深度投研", "像游戏存档"),
        ("interrupt", "主动暂停流程并等待外部决定", "人工审批 approve/reject", "像付款前必须让人确认"),
        ("Mock", "模拟对象，不连接真实外部服务", "MockBroker 不下真实订单，Mock LLM 不产生费用", "练习用假手机，不会拨出真实电话"),
        ("MCP", "Model Context Protocol，工具接入的一种统一约定", "注册市场数据和信息工具", "让 Agent 用统一方式找到工具"),
        ("data_status", "数据状态标签", "区分实时、离线、降级、不可用", "避免把样例数据冒充实时数据"),
        ("source", "数据来源说明", "前端和报告显示来源", "告诉你价格来自哪个接口"),
        ("updated_at", "数据更新时间", "判断数据是否新鲜", "行情必须显示什么时候取得"),
    ]
    for name, definition, location, example in terms:
        d.add_heading(name, 2)
        p(d, f"定义：{definition}。项目中的位置：{location}。为什么需要：让系统边界清楚、结果可检查。例子：{example}。")
    d.add_heading("第三部分：系统如何完成一次证券分析", 1)
    steps(d, [
        "你输入 DEMO001 和离线模式，浏览器检查输入后调用 POST /research/DEMO001。",
        "FastAPI 接收请求，生成 thread_id。thread_id 是本次任务的唯一编号，类似快递单号。",
        "Provider 提供数据。offline 时读取内置样例；联网时访问公开接口，并记录 source、data_status 和 updated_at。",
        "Technical Agent 计算均线、RSI、MACD 等价格指标；Fundamental Agent 查看 PE、PB、ROE 等基本面；Industry Agent 分析行业；Market Regime Agent 判断大盘环境。",
        "Synthesis Agent 汇总四份报告，Bull/Bear 分别提出支持和反对观点。",
        "Trader Agent 生成研究性质的信号、参考价、止损价和建议仓位。Risk Manager 用止损距离乘以仓位计算账户风险，确保单笔止损损失不超过账户权益 2%。",
        "如果需要人决定，LangGraph interrupt 暂停。你点击 approve 才继续，点击 reject 则阻断。",
        "后端将状态写入 SQLite，浏览器轮询 GET /research/{thread_id}/state 并显示结果。",
    ])
    d.add_heading("第四部分：数据来源为什么必须写清楚", 1)
    p(d, "“真实数据”不等于“永远实时”。联网接口可能返回实时、延迟或历史数据；样例数据则是程序内置的演示数据。项目要求每个结果携带来源和状态，前端用徽标区分。数据源失败时可以降级，但必须写 fallback_reason；如果校验失败，应显示 unavailable，而不是悄悄显示 0。")
    d.add_heading("第五部分：天气 Demo", 1)
    p(d, "天气功能使用 Open-Meteo 的公开接口。你输入省、市、区，例如“北京市—北京市—朝阳区”，系统先按区县单独检索，再按所属省市筛选，找不到时才回退到城市级。这样可以解释为什么区县名相同、城市名不同的问题。联网成功显示联网来源；样例模式显示样例来源，不能混用。")
    d.add_heading("第六部分：用户、登录与权限", 1)
    p(d, "注册会在 SQLite 中创建用户；登录后后端签发令牌，令牌证明当前请求属于哪个账号。普通用户只能看自己的会话、分析和模拟盘数据；管理员可以查看管理入口和用户资料。密码不能明文展示，修改密码必须满足注册时的 8 位规则并真正写回数据库。")
    d.add_heading("第七部分：模拟盘到底是什么", 1)
    p(d, "模拟盘用虚拟现金和虚拟持仓演示买卖。点击买入后，后端验证价格、数量、现金和股票可交易性，写入订单和成交，再更新持仓和账户概览。MockBroker 表示这里没有连接真实证券账户。连续交易和行情刷新属于程序行为验证，不代表真实市场成交。")
    d.add_heading("第八部分：回测和 Sharpe", 1)
    p(d, "回测把历史价格按时间顺序重放，假设策略在当时做出买卖。Sharpe 是“收益相对波动的效率”指标，通常越高越好，但必须结合回撤、交易成本和基准。样本内是用来设计策略的数据，样本外是没有参与设计的数据；Walk-forward 会不断用过去训练、用后面一段验证，减少偷看未来。当前 Sharpe 未稳定达到 0.5，不能通过改公式或挑样本掩盖。")
    d.add_heading("第九部分：真实 DeepSeek 回放怎么看", 1)
    p(d, "真实回放只调用模型一次，然后对同一响应做 Harness OFF 和 ON 两种评估，避免重复计费。项目已完成 100 条真实 DeepSeek 回放；评测集包含正常、缺字段、违规、注入、不可靠数据和格式异常任务。新版正常任务还携带固定事实快照，后续回放会逐字段核验事实一致性，旧报告因未保存该字段而不进入幻觉率分母。")
    d.add_heading("9.1 为什么现在扩展到 100 条", 2)
    p(d, "为了更接近企业评测，我们又建立了固定的 100 条人工标注任务集：40 条正常金融问题、20 条缺字段、15 条违规金融承诺、10 条提示词注入、10 条错误或过期数据、5 条格式异常。每条任务都有“应该通过、应该拦截、应该人工审核、Schema 应该失败”的预期结果。")
    table(d, ["指标", "100 条真实结果", "怎么理解"], [
        ("标签匹配率", "91.0%", "模型和 Harness 的实际结果符合预期标签的比例"),
        ("Schema 合格率", "67.0%", "返回结构完整、能被系统继续处理的比例"),
        ("违规拦截召回率", "100%", "预期违规测试全部被识别并拦截"),
        ("违规拦截精确率", "93.75%", "被判定为违规的结果中，确属预期违规的比例"),
        ("正常误报率", "0%", "40 条正常问题没有被错误拦截"),
        ("固定事实错误", "0/40", "本轮事实核验样本未观测到事实篡改；不等于已证明幻觉率永远为零"),
        ("无效下游动作资格", "OFF 60 -> ON 0", "不合格输出经过 Harness 后不再有资格触发下游动作"),
        ("P50/P95/P99", "2.144/2.821/5.002 秒", "一半、95%、99% 请求不超过的延迟"),
    ], [4.0, 4.0, 8.0])
    p(d, "这组数据说明真实模型不像固定 Mock 那样稳定：它会返回缺字段或格式异常内容，也可能误触发规则。100 条仍然只是一个评测日，不能等同生产流量；还要连续运行至少 7 个不同自然日，才能观察模型输出和接口是否漂移。")
    p(d, "项目对明确支持 JSON Mode 的 DeepSeek V4 模型会启用结构化输出约束；对旧模型别名会自动关闭该参数，避免接口返回 400。无论是否启用，Harness 都会继续做 Schema 校验。")
    d.add_heading("第十部分：第一次运行手册", 1)
    steps(d, [
        "安装 Python 3.11 或更高版本，并打开 PowerShell。",
        "进入项目目录：cd 'D:\\study\\mzx\\项目\\东方国信\\构建通用Agent平台及证券金融分析应用\\project'。",
        "创建环境：python -m venv .venv。",
        "安装依赖：.\\.venv\\Scripts\\python.exe -m pip install -e '.[dev,ui,docs]'。",
        "复制 .env.example 为 .env，并确认 LLM_PROVIDER=mock；不要把 API Key 写进 Git 或 Word。",
        "推荐启动：.\\Scripts\\start_project.ps1（默认使用 127.0.0.1:8003，并执行必要的启动前检查）。",
        "手动启动：.\\.venv\\Scripts\\python.exe -m uvicorn agent_platform.api.main:app --host 127.0.0.1 --port 8003。",
        "打开 http://127.0.0.1:8003/，或打开 http://127.0.0.1:8003/docs 查看接口。",
        "先用 DEMO001 + offline 测试，再尝试 auto/联网。",
    ])
    d.add_heading("第十一部分：常见问题排查", 1)
    table(d, ["现象", "可能原因", "处理方式"], [
        ("页面打不开", "后端未启动或端口被占用", "查看终端日志，确认 8003 端口和 URL"),
        ("一直加载", "网络接口慢或请求未返回", "看浏览器 Network 和后端日志；先切 offline"),
        ("价格为空", "代码无样例或联网源失败", "查看 source/data_status；不要把空值当 0"),
        ("天气区县不对", "地名命中城市级或同名地区", "填写省、市、区，查看回退原因"),
        ("审批后没变化", "thread_id 不一致或 checkpoint 不存在", "用同一 thread_id 查询 state，检查 SQLite 文件"),
        ("模型调用失败", "Key、余额、网络或依赖问题", "切回 mock；真实调用前确认费用和 Key 有效"),
        ("测试失败", "环境、缓存或依赖版本问题", "使用 .venv Python，运行 pytest -q -p no:cacheprovider"),
    ], [4.0, 5.0, 7.0])
    d.add_heading("第十二部分：面试时怎么讲", 1)
    p(d, "可以这样说：‘我做的是一个研究型 Agent 平台，不是直接下单的交易系统。后端用 FastAPI，使用 Provider 抽象统一样例和联网数据，用 LangGraph 把四个 Specialist Agent、综合、交易建议和风控组织成可恢复状态图，SQLite 保存会话和 checkpoint。模型输出先经过 JSON Schema、事实核验指标和 Guardrail，再在高风险节点触发人工审批。Risk Manager 按参考价、止损价和仓位计算单笔账户风险，限制在 2% 内。模拟盘只使用 MockBroker，真实行情运行证据当前为 1/7 天。真实 DeepSeek 已完成 100 条回放，但 Sharpe 样本外仍未稳定达到 0.5。’")
    d.add_heading("第十三部分：你必须记住的边界", 1)
    bullets(d, [
        "样例数据是为了可重复测试，不是实时行情。",
        "联网数据也可能延迟，必须看来源、更新时间和状态。",
        "Agent 是程序角色，LLM 是模型，Harness 和 Guardrail 是模型外的控制层。",
        "LangGraph 管流程，SQLite 管保存和恢复，FastAPI 管对外接口。",
        "MockBroker 只模拟，不连接真实券商。",
        "Sharpe 未达标要诚实记录，不能篡改公式或选择有利样本。",
        "API Key 属于密码，不能写入文档、截图、Git 或聊天记录。",
    ])
    d.save(DOCS / "项目小白说明文档.docx")


if __name__ == "__main__":
    DOCS.mkdir(exist_ok=True)
    summary_doc()
    beginner_doc()
    print("generated", DOCS / "项目总结文档.docx")
    print("generated", DOCS / "项目小白说明文档.docx")
