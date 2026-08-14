"""生成项目交付文档；内容以当前代码和原始任务书为准。"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FONT = "Calibri"
CN_FONT = "微软雅黑"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)


def font(run, size=11, bold=False, color=None) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), CN_FONT)


def setup(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.text = "通用 Agent 平台及证券金融分析应用 | 项目交付资料"
    for run in header.runs:
        font(run, 9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("东方国信实习项目 | 仅供研究参考")
    for run in footer.runs:
        font(run, 9, color=GRAY)


def title(doc: Document, value: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(value)
    font(r, 23, True, RGBColor(11, 37, 69))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(subtitle)
    font(r, 12, color=GRAY)


def h(doc: Document, text: str, level=1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.25)
    for run in p.runs:
        font(run)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    for run in p.runs:
        font(run)


def step(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    for run in p.runs:
        font(run)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    widths = [9360 // len(headers)] * len(headers)
    widths[-1] += 9360 - sum(widths)
    for i, cell in enumerate(t.rows[0].cells):
        cell.width = Inches(widths[i] / 1440)
        cell.text = headers[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tcpr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8EEF5")
        tcpr.append(shd)
        for run in cell.paragraphs[0].runs:
            font(run, 10, True, DARK)
    for values in rows:
        cells = t.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = Inches(widths[i] / 1440)
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    font(run, 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def save(doc: Document, name: str) -> None:
    path = DOCS / name
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(path)


def make_summary() -> None:
    d = Document(); setup(d); title(d, "通用 Agent 平台及证券金融分析应用", "项目总结文档 | 面向导师与技术面试的事实型说明")
    h(d, "一、执行摘要")
    para(d, "本项目以证券金融分析为复杂业务场景，验证 Harness Engineering、Loop Engineering 和 Graph Engineering 的组合落地。当前系统不是券商生产交易系统，而是一套可运行、可测试、可审计的研究型平台原型：联网模式可通过 MCP 访问公开行情，离线模式使用同一工作流与确定性样例数据；交易侧只使用本地 MockBroker，任何交易建议必须经过风控和人工确认。")
    table(d, ["维度", "当前事实"], [["编排", "LangGraph StateGraph；四路 Specialist 并行，支持条件路由、interrupt/resume、SQLite checkpoint"], ["Agent", "技术、基本面、行业、市场状态、综合、交易、风控等节点；Specialist 内部接入 AgentLoop + AgentHarness"], ["数据", "MCP Registry 统一封装在线/离线行情与财务工具；auto 失败可降级并标记来源"], ["安全", "Source / Schema / 关键词 / 交叉验证 / 人工审批 / MockBroker 隔离"], ["未达标", "回测 Sharpe 仍低于说明书阈值，不能宣称策略有效或可盈利"]])
    h(d, "二、原始说明书目标与实现映射")
    table(d, ["说明书要求", "实现位置", "状态"], [["Harness 九大组件", "SPEC.md、Rule、Skill、Workflow、Scripts、MCP、SubAgents、AGENTS.md、看板文件", "已落地"], ["Loop 五要素与持久化记忆", "core/agent_loop.py、core/loop_memory.py、MemoryScope、Echo Demo", "已落地"], ["Graph 并行、条件边、checkpoint", "finance/securities_graph.py、Workflow/securities_analysis.workflow.json", "已落地"], ["四 Specialist 结构化输出", "specialist_runtime.py + 四个金融 Agent", "已落地"], ["Bull/Bear 与人工澄清", "bull_bear_debate.py、debate_approval interrupt", "已落地"], ["20 只证券端到端", "Scripts/validate_deliverables.py", "可运行；execute/no_trade 比例取决于输入数据"], ["Sharpe > 0.5", "回测模块", "未达标，需继续研究"]])
    h(d, "三、系统架构")
    para(d, "展示层使用 frontend_prototype.html；接口层使用 FastAPI；ApplicationService 负责业务门面和 LangGraph 生命周期；LangGraph 负责跨 Agent 拓扑；每个 Specialist 通过 AgentLoop 形成规划、工具调用、观察、反思、决策闭环，再由 AgentHarness 执行结构、来源和安全护栏；MCP Registry 是行情与财务工具的统一入口；SQLite 保存会话、Loop 记忆、模拟账户和 LangGraph checkpoint。")
    h(d, "四、深度投研主流程", 2)
    for text in ["客户端提交 symbol 与 data_mode，服务生成 thread_id。", "四个 Specialist 并行取数和计算技术/基本面/行业/市场状态结果，并追加审计记录。", "Synthesis 汇总结果和 Bull/Bear 证据；置信度不高于 0.30 进入 no_trade。", "较高置信度且辩论检查阻断时进入 debate_approval；人工批准后继续，拒绝则 block。", "Trader 生成信号；仓位超过 10% 触发 HumanApprovalRequired，LangGraph 中断并写入 checkpoint。", "Risk Manager 和 TradingHarness 继续执行风控与 Pre-Flight，最终输出 execute、manual_review 或 block。"]: step(d, text)
    h(d, "五、工程质量与验证")
    bullet(d, "离线验证强制使用 LLM_PROVIDER=mock、MARKET_DATA_PROVIDER=sample；MCP 注册表在工具函数执行前阻断网络工具。")
    bullet(d, "所有金融结果包含 source、updated_at、data_status、fallback_reason 和免责声明；不可用时返回显式错误。")
    bullet(d, "SQLite checkpoint、LoopMemory 和 PaperBroker 均支持重建后恢复；前端模拟盘不再以浏览器内存作为账户事实源。")
    bullet(d, "全量 pytest 当前通过，保留一个 httpx/starlette 依赖弃用警告；该警告不代表业务测试失败。")
    h(d, "六、边界、风险与后续计划")
    para(d, "DCF 当前为 earnings-to-FCFF proxy：缺少完整 EBIT、税项、折旧摊销、资本开支和营运资本变动，因此结果以 low confidence、limitations 和 warnings 明确标注，不应称为完整财报 DCF。回测 Sharpe 未达到说明书阈值，不能通过文档美化；后续应补充真实历史财务字段、Walk-Forward 验证、多因子组合和样本外检验。")
    bullet(d, "当前定位：研究级、可运行、可扩展的工程原型；不是生产券商交易系统。")
    bullet(d, "上线前还需身份认证、权限管理、密钥托管、审计合规、监控告警、数据质量 SLA 和人工审批操作台。")
    save(d, "项目总结文档.docx")


def make_beginner() -> None:
    d = Document(); setup(d); title(d, "项目小白说明文档", "从零理解 Harness、LangGraph、FastAPI 与证券分析平台")
    h(d, "一、先记住一句话")
    para(d, "这个项目不是一个会自动炒股的机器人，而是一条“数据进入 - 多个分析员分别研究 - 汇总判断 - 风控拦截 - 人工确认 - 本地模拟”的流水线。每一步都有输入、输出、来源和失败处理。")
    h(d, "二、三个最容易混淆的框架")
    table(d, ["名词", "白话解释", "在本项目中的作用"], [["Agent", "能调用工具、根据目标完成任务的程序角色", "技术面、基本面、行业面、市场状态分别研究"], ["Loop", "一个 Agent 内部的工作循环：规划→调用工具→看结果→反思→决定继续或结束", "保证 Specialist 不是只调用一次函数就结束"], ["Graph / LangGraph", "把多个 Agent 画成有向流程图并按依赖执行的编排框架", "四路并行，之后汇总、路由、审批和风控"], ["Harness", "包住 Agent 的安全外壳，像飞机的检查单和黑匣子", "校验输入输出、来源、违规词、重试、审计、人工审批"]])
    h(d, "三、FastAPI 是什么")
    para(d, "FastAPI 是 Python 的 Web 接口框架。它把 Python 函数变成浏览器或前端可以调用的 HTTP 地址。例如 POST /research/DEMO001 表示启动一次深度投研，GET /research/{thread_id}/state 表示查询状态。Pydantic 模型负责检查请求字段，FastAPI 自动生成 /docs 接口页面。")
    h(d, "四、一次深度投研到底发生什么")
    for text in ["前端发送股票代码和模式。auto 代表允许联网并可降级；offline 代表零网络。", "MCP Registry 根据工具名取行情。MCP 是统一的工具插座，工具返回固定信封：ok、data、source、updated_at、error。", "技术 Agent 计算 MA、MACD、RSI 等；基本面 Agent 读取 PE、PB、ROE、资产负债率和 DCF 代理；行业 Agent 研究行业；市场状态 Agent 判断大盘。", "四个结果汇总进 Synthesis。置信度 ≤30% 直接 no_trade，不生成交易信号。", "交易信号仓位超过 10% 会暂停，用户 approve/reject 后 LangGraph 从 SQLite checkpoint 恢复。", "风控和 Pre-Flight 最后决定 execute、manual_review 或 block；execute 也只是研究建议，不会自动下真实订单。"]: step(d, text)
    h(d, "五、底层数据层是什么")
    para(d, "数据层不是一个抽象口号，而是负责“从哪里取得数据、把数据变成什么格式、失败时怎么办”。SampleMarketDataProvider 读取项目 data 目录里的确定性 CSV，只用于离线测试；MCPMarketDataProvider 是统一适配器，在线和离线都经过 MCP Registry；AkShare 工具访问公开数据，失败时返回明确错误，不能偷偷编造价格。")
    h(d, "六、必须理解的术语")
    table(d, ["术语", "解释"], [["LLM", "大语言模型，负责理解和生成文本；本项目测试默认用 mock，避免真实 API 费用"], ["Schema", "数据格式合同，规定必须有哪些字段、字段类型是什么"], ["Guardrail", "护栏，输出不符合规则时阻断"], ["MCP", "统一调用外部工具的协议/注册层，隐藏 AkShare 等具体实现"], ["checkpoint", "工作流断点快照，中断后可以从原位置继续"], ["thread_id", "一次投研任务的唯一编号，用来查询和恢复状态"], ["MockBroker", "本地模拟撮合器，只改本地现金和持仓，不连接券商"], ["Sharpe", "收益相对波动风险的指标；本项目策略结果未达到说明书阈值"]])
    h(d, "七、你在浏览器中看到什么")
    bullet(d, "/docs：FastAPI 自动接口文档。")
    bullet(d, "证券分析：查看请求区间、K 线和指标，状态徽标区分实时、离线、降级、不可用。")
    bullet(d, "深度投研：观察 11 个工作流节点、四个 Agent 卡片和审批面板。")
    bullet(d, "模拟盘：账户由后端 SQLite + MockBroker 管理，订单和持仓重启后仍可恢复。")
    h(d, "八、不要误解的地方")
    para(d, "联网成功不等于数据一定是此刻交易所撮合价；公开数据源有延迟、接口限制和更新时间。离线模式显示的是样例数据，不是真实行情。DCF 是低可信估值代理，Sharpe 未达标，系统不能据此承诺收益或自动交易。")
    save(d, "项目小白说明文档.docx")


def make_guide() -> None:
    d = Document(); setup(d); title(d, "项目指导说明书", "从空目录手把手搭建可运行的工程原型")
    h(d, "一、准备环境")
    for text in ["安装 Python 3.11，并确认 python --version。", "创建虚拟环境：python -m venv .venv；Windows 激活：.venv\\Scripts\\activate。", "安装项目依赖：python -m pip install -e .；需要联网行情时再安装项目定义的 AkShare 额外依赖。", "复制 .env.example 为 .env，仅在本地填写密钥；严禁把 .env 提交到 Git。"]: step(d, text)
    h(d, "二、先搭通用平台")
    table(d, ["顺序", "创建内容", "验收方式"], [["1", "Rule、Skill、SPEC、AGENTS.md 和 checklist", "任何新开发者能找到约束、工具和成功标准"], ["2", "ToolRegistry 与 AgentLoop", "Echo Agent 能完成规划、调用、观察、反思、决策"], ["3", "LoopMemory SQLite", "新建对象重新打开同一数据库仍能读到记录"], ["4", "AgentHarness 与 Guardrail", "缺 source、非法 Schema、违规词会被阻断"], ["5", "TaskStateNamespace", "不同 task_id 写入不同目录，../escape 被拒绝"]])
    h(d, "三、实现 LangGraph 工作流")
    for text in ["定义 TypedDict 状态；并行节点写不同 key，列表字段明确 reducer。", "创建 StateGraph，START 扇出到 technical、fundamental、industry、market_regime。", "四路都指向 synthesis，LangGraph 会等待所有前驱完成。", "使用 add_conditional_edges 定义 no_trade、debate_approval、trader_agent 等分支。", "使用 interrupt() 暂停，并为生产模式注入 SqliteSaver；resume 必须携带同一 thread_id。", "在 Workflow/securities_analysis.workflow.json 同步节点、边、状态键和中断声明，并用一致性测试防漂移。"]: step(d, text)
    h(d, "四、接入真实公开数据")
    para(d, "先用 offline 跑通全部功能，再切 auto。在线工具必须在 MCP Registry 注册，requires_network=True；离线工具必须独立注册，函数体不能 import 网络库。每个成功信封带 source 和 updated_at，失败信封 data 必须为 None。使用 000001 或 600519 等真实代码测试，不要用 DEMO001 验证联网。")
    h(d, "五、实现四类金融 Agent")
    for text in ["技术：取历史 OHLCV，由确定性指标函数计算 MA/MACD/RSI/KDJ/ATR/CCI。", "基本面：取估值和财务指标；缺少完整报表时 DCF 必须标记 proxy、low confidence 和 limitations。", "行业：输出行业、景气、资金流和龙头证据。", "市场状态：输出指数趋势、风险偏好和 regime。", "每个 Agent 通过 SpecialistRuntime 进入 AgentLoop，再由 AgentHarness 做 Schema、source 和 CrossValidator 校验。"]: bullet(d, text)
    h(d, "六、接入交易决策和模拟盘")
    for text in ["Synthesis 输出 signal 和 confidence；低于阈值 no_trade。", "Trader 计算仓位；超过 10% 抛出 HumanApprovalRequired。", "Risk Manager 校验交易信号；TradingHarness 运行 Pre-Flight。", "模拟盘只使用 MockBroker；通过 /paper-trading/accounts 创建账户、/orders 下单、/refresh 刷新行情，SQLite 保存快照。", "连续策略模拟使用 POST /paper-trading/runs，结果通过 GET /paper-trading/runs/{run_id} 恢复。"]: step(d, text)
    h(d, "七、运行与验收命令")
    for text in ["设置离线测试变量：$env:LLM_PROVIDER='mock'; $env:MARKET_DATA_PROVIDER='sample'。", "运行服务：.venv\\Scripts\\python.exe -m uvicorn agent_platform.api.main:app --app-dir . --host 127.0.0.1 --port 8003。", "打开 http://127.0.0.1:8003/docs，先调用 GET /health。", "调用 POST /research/DEMO001?data_mode=offline，查询 state，若 interrupted 再调用 resume。", "调用 POST /paper-trading/runs，symbols=[DEMO001]、data_mode=offline、days=20。", "运行全量测试：.venv\\Scripts\\python.exe -m pytest -q -p no:cacheprovider。", "联网验收：切换 data_mode=auto，使用真实 A 股代码，检查 source、updated_at 和 data_status；严禁把结果用于真实下单。"]: step(d, text)
    h(d, "八、交付前检查清单")
    for text in ["pytest、compileall、前端 node --check 全部通过。", "offline 模式零网络调用；auto 失败时 fallback_reason 非空。", "所有输出含来源、时间和免责声明。", "人工审批 reject 不会进入风险或模拟执行节点。", "Sharpe 若未达阈值，报告必须标记未达标。", "检查 git status，确认 .env、密钥、数据库和缓存未提交。"]: bullet(d, text)
    save(d, "指导说明书.docx")


def make_work_report() -> None:
    source = DOCS / "_source" / "internship_template.docx"
    d = Document(source)
    d.paragraphs[0].text = "工作总结报告"
    for run in d.paragraphs[0].runs:
        font(run, 18, True, RGBColor(0, 0, 0))
    t = d.tables[0]
    def set_cell(cell, value: str, size=10):
        cell.text = value
        for p in cell.paragraphs:
            for run in p.runs:
                font(run, size, False)
    set_cell(t.cell(0, 0), "学生姓名")
    set_cell(t.cell(0, 1), "米志轩")
    set_cell(t.cell(0, 2), "项目及岗位名称")
    set_cell(t.cell(0, 3), "东方国信-通用 Agent 平台及证券金融分析应用")
    set_cell(t.cell(1, 0), "项目起止时间：2026/7/20-2026/8/10")
    set_cell(t.cell(2, 0), "每周工作总结")
    weeks = """第一周（7月20日-7月22日）\n1. 工作内容：阅读原始任务书，理解 Prompt、Context、Loop、Graph、Harness 的工程关系；搭建 Python 项目结构、规则目录、Workflow、Scripts、MCP 和测试基础设施。\n2. 完成情况：实现 AgentHarness、Guardrail、ToolRegistry、AgentLoop 和基础金融分析模块；完成离线样例与 FastAPI 接口。\n3. 困难与收获：理解可靠 Agent 不能只靠 Prompt，必须有状态、工具、验证和审计。\n\n第二周（7月29日-8月2日）\n1. 工作内容：迁移证券主链到 LangGraph，完成四路 Specialist 并行、Synthesis、Trader、Risk Manager、TradingHarness、interrupt/resume 与 SQLite checkpoint。\n2. 完成情况：补齐 workflow JSON 与真实图一致性测试；实现低置信 no_trade、人工审批和服务重建恢复。\n3. 困难与收获：解决中断状态保存、并行状态 reducer、前后端字段契约和真实/离线数据状态区分问题。\n\n第三周（8月3日-8月10日）\n1. 工作内容：将 Specialist 接入 AgentLoop + AgentHarness，补齐 working/project/organization 三级记忆、任务目录隔离、Echo Demo、MCP 统一行情入口和持久化模拟盘 API。\n2. 完成情况：前端模拟盘改为以后端 SQLite + MockBroker 为事实源；新增连续多交易日模拟 API；完成 DCF 口径修正并明确 proxy 边界。\n3. 验收结论：工程主线已完成并通过全量测试；回测 Sharpe 尚未达到说明书阈值，作为后续策略研究任务保留。"""
    set_cell(t.cell(3, 0), weeks, 9)
    set_cell(t.cell(4, 0), "工作收获")
    set_cell(t.cell(5, 0), "掌握了 LangGraph StateGraph、SQLite checkpoint、FastAPI REST API、MCP 工具注册、AgentLoop 五要素、Harness Guardrail、金融指标计算、MockBroker 模拟撮合和 pytest 集成测试。更重要的是理解了数据来源、状态恢复、异常处理和人工审批在金融系统中的工程意义。")
    set_cell(t.cell(5, 2), "项目中所使用的工具")
    set_cell(t.cell(5, 3), "Python 3.11、FastAPI、LangGraph、SQLite、pandas、pytest、AkShare、MCP Registry、MockBroker、HTML/JavaScript、VS Code。")
    set_cell(t.cell(6, 0), "请列举工作过程中的其他收获")
    set_cell(t.cell(6, 2), "知识：理解 Harness/Loop/Graph 分层。技能：能设计状态契约、离线测试和持久化恢复。职业素养：形成先读说明书、再改代码、最后用测试证明的工作习惯。")
    set_cell(t.cell(7, 0), "参与后是否会将此类工作为未来求职目标，为什么？")
    set_cell(t.cell(7, 2), "有意向。该项目同时训练了 Agent 编排、后端 API、数据工程、测试和金融业务理解能力。后续希望继续提升真实 LLM 接入、分布式执行、数据质量与量化策略研究能力。")
    set_cell(t.cell(8, 0), "身边的导师是否有让你觉得佩服的，他们有哪些特点？")
    set_cell(t.cell(8, 2), "导师提供了结构清晰、验收标准明确的任务书，帮助我把抽象的 Agent 概念转化为可执行的工程任务。令我佩服的是对边界条件、可复现性和项目交付质量的重视。")
    set_cell(t.cell(9, 0), "自己还有哪些方面需要提升？")
    set_cell(t.cell(9, 2), "需要提升真实大模型评测与提示/工具调用策略、分布式 LangGraph 部署、金融报表口径、组合优化、Walk-Forward 回测和生产级安全合规能力。Sharpe 未达标也说明策略研究不能只完成工程链路，还需要持续实验。")
    set_cell(t.cell(10, 0), "请客观陈述一下自己的项目感受")
    set_cell(t.cell(10, 2), "项目让我认识到，Agent 项目的难点不只是让模型输出文字，而是让系统在数据缺失、网络失败、人工中断和服务重启后仍然可控、可解释、可恢复。当前成果是研究级工程原型，不是生产交易系统；我对已完成能力和未达标指标均作了如实记录。")
    save(d, "米志轩-通用Agent平台及证券金融分析应用-工作总结(7.20-8.10).docx")


if __name__ == "__main__":
    make_summary()
    make_beginner()
    make_guide()
    make_work_report()
